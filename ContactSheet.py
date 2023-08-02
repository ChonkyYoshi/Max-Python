from pathlib import Path
from zipfile import ZipFile
from regex import match, search
from PIL import Image
from imagequant import quantize_pil_image
import docx
from shutil import copyfile, rmtree


def ExtractImages(File: Path):

    zipfile = ZipFile(File)
    TempDir = Path(File.parent.as_posix() + '/Temp')
    TempDir.mkdir(exist_ok=True)
    for file in zipfile.namelist():
        if match(r'.*?/media/.*?\.(jpeg|jpg|png|emf|wmf|wdp)', file):
            zipfile.extract(file, TempDir)
    if File.suffix == '.pptx' or File.suffix == '.story':
        for rel in zipfile.namelist():
            if match(r'.*?/slides/_rels', rel):
                zipfile.extract(rel, TempDir)
    return TempDir


def CleanTempDir(Tempdir: Path, compress=True):

    for index, file in enumerate(list(Tempdir.glob('*/media/*'))):
        i = len(list(Tempdir.glob('*/media/*')))
        yield f'Converting all to png, image {index + 1} of {i}'
        if not file.suffix == '.png':
            with Image.open(file.as_posix()) as im:
                try:
                    im.save(f'{file.parent.as_posix()}/{file.stem}.png')
                    file.unlink(missing_ok=True)
                except OSError:
                    continue
    if compress:
        for index, file in enumerate(list(Tempdir.glob('*/media/*'))):
            i = len(list(Tempdir.glob('*/media/*')))
            yield f'Compressing, image {index + 1} of {i}'
            if file.suffix == '.png':
                with Image.open(file.as_posix()) as im:
                    new_im = quantize_pil_image(im)
                    new_im.save(file.as_posix()[:-3] + 'png')


def FillCS(TempDir: Path, File: Path):

    CSPath = Path(File.parent.as_posix() + '/Contact Sheets')
    CSPath.mkdir(exist_ok=True)
    CS = docx.Document()

    i = len(list(TempDir.glob('*/media/*')))
    for index, pic in enumerate(list(TempDir.glob('*/media/*'))):
        if pic.is_dir() or pic.suffix == '.rels':
            continue
        Table = CS.add_table(rows=5, cols=2, style='Table Grid')
        Table.cell(0, 0).merge(Table.cell(0, 1)).text = pic.name
        Table.cell(1, 0).merge(Table.cell(1, 1))
        Table.cell(2, 0).merge(Table.cell(2, 1))
        Table.cell(3, 0).text = 'Source'
        Table.cell(3, 1).text = 'Target'
        try:
            yield f'Filling in Contact Sheet, image {index + 1} of {i}'
            run = Table.cell(1, 0).paragraphs[0].add_run()
            run.add_picture(pic.as_posix(),
                            width=(CS.sections[0].page_width -
                                   (CS.sections[0].right_margin +
                                    CS.sections[0].left_margin)))
        except Exception:
            run = Table.cell(1, 0).paragraphs[0].add_run()
            run.text = f'''An error occured while trying to insert
the image, please check the Error folder and manually insert the image.
Name of the image: {pic.name}'''
            ErrorDir = Path(File.parent.as_posix() +
                            '/Errors_' + File.stem)
            ErrorDir.mkdir(exist_ok=True)
            copyfile(pic.as_posix(),
                     ErrorDir.as_posix() + '/' + pic.name)
        if File.suffix == '.pptx' or File.suffix == '.story':
            locations = LocateImage(TempDir, pic.stem)
            if len(locations) == 0:
                Table.cell(2, 0).add_paragraph(
                    'only present in Master Slide', style='List Bullet')
                CS.add_page_break()
            else:
                for location in locations:
                    Table.cell(2, 0).add_paragraph(location,
                                                   style='List Bullet')
                CS.add_page_break()
        else:
            CS.add_page_break()
    CS.save(f'{CSPath.as_posix()}/CS_{File.name}.docx')
    rmtree(TempDir)


def LocateImage(TempDir: Path, ImageName: str):

    Locations = []
    for entry in TempDir.rglob('*.rels'):
        with open(entry, 'r') as File:
            rel = File.read()
        if search(ImageName + r'\.', rel):
            match = search(r'(\w+)\.', entry.name)
            Locations.append(entry.
                             name[match.start():match.end()])  # type: ignore
    return Locations
