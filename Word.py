from helper import CopyParFormatting, CopyRunFormatting, isolate_run
import docx
from pathlib import Path
from regex import match


def BilTable(File: Path):

    li = list()
    doc = docx.Document(File)
    for index, table in enumerate(doc.tables):
        yield f'processing table {index + 1} of {len(doc.tables)}'
        for c in range(len(table.columns)):
            for r in range(len(table.rows)):
                if table.cell(r, c)._tc not in li:
                    li.append(table.cell(r, c)._tc)
                    for par in table.cell(r, c).paragraphs:
                        prevpar = par.insert_paragraph_before()
                        CopyParFormatting(prevpar, par)
                        for run in par.runs:
                            prevrun = prevpar.add_run()
                            prevrun.text = run.text
                            CopyRunFormatting(prevrun, run)
                            prevrun.font.hidden = True
        li.clear()
    i = len(doc.paragraphs)
    for index, par in enumerate(doc.paragraphs):
        yield f'processing paragraph {index + 1} of {i}'
        table = doc.add_table(rows=1, cols=2)
        par._p.addnext(table._tbl)
        SPar = table.cell(0, 0).paragraphs[0]
        TPar = table.cell(0, 1).paragraphs[0]
        CopyParFormatting(SPar, par)
        CopyParFormatting(TPar, par)
        for run in par.runs:
            SRun = SPar.add_run()
            TRun = TPar.add_run()
            SRun.text = run.text
            TRun.text = run.text
            CopyRunFormatting(SRun, run)
            CopyRunFormatting(TRun, run)
        par._element.getparent().remove(par._element)
    doc.save(File.parent.as_posix() + '/Bil_' + File.name)


def Doc2PDF(WordApp, File: Path,
            ARev: bool = False,
            DRev: bool = False,
            Com: bool = False,
            Overwrite: bool = False):

    Doc = WordApp.Documents.Open(File.as_posix(), Visible=False)
    PdfDir = Path(File.as_posix() + '/PDF')
    PdfDir.mkdir(exist_ok=True)
    if Doc.Revisions.Count > 0 and ARev:
        Doc.AcceptAllRevisions()
    if Doc.Revisions.Count > 0 and DRev:
        Doc.RejectAllRevisions()
    if Doc.Comments.Count > 0 and Com:
        Doc.DeleteAllComments()
    if Overwrite:
        Doc.Save()
    Doc.SaveAs2(f'{PdfDir.as_posix()}/{File.name}.pdf', FileFormat=17)
    Doc.Close()


def AcceptRevisions(WordApp, File: Path,
                    ARev: bool = False,
                    DRev: bool = False,
                    Com: bool = False,
                    Overwrite: bool = False):

    Doc = WordApp.Documents.Open(File.as_posix(), Visible=False)

    if Doc.Revisions.Count > 0 and ARev:
        Doc.AcceptAllRevisions()
    if Doc.Revisions.Count > 0 and DRev:
        Doc.RejectAllRevisions()
    if Doc.Comments.Count > 0 and Com:
        Doc.DeleteAllComments()
    if Overwrite:
        Doc.Save()
    else:
        match File.suffix:
            case '.docx':
                Doc.SaveAs2(File.parent.as_posix() + 'NoRev_' +
                            File.stem, FileFormat=12)
            case '.docm':
                Doc.SaveAs2(File.parent.as_posix() + 'NoRev_' +
                            File.stem, FileFormat=13)
            case '.doc':
                Doc.SaveAs2(File.parent.as_posix() + 'NoRev_' +
                            File.stem, FileFormat=0)
    Doc.Close()


def PrepStoryExport(File: Path, Regex: str = ''):

    Doc = docx.Document(File)
    for index, par in enumerate(Doc.paragraphs):
        yield f'paragraph {index + 1} of {len(Doc.paragraphs)}'
        for run in par.runs:
            run.font.hidden = True
    for table in Doc.tables:
        for index, SkipCol in enumerate(table.columns):
            yield f'column {index + 1} of {len(table.columns)}'
            if not index == 3:
                for cell in SkipCol.cells:
                    for par in cell.paragraphs:
                        for run in par.runs:
                            run.font.hidden = True
            else:
                for cell in SkipCol.cells:
                    for par in cell.paragraphs:
                        for run in par.runs:
                            if match(Regex, run.text) and Regex != '':
                                start = match(Regex,
                                              run.text).start()  # type: ignore
                                end = match(Regex,
                                            run.text).end()  # type: ignore
                                hidden_run = isolate_run(par, start, end)
                                hidden_run.font.hidden = True
                for par in table.cell(0, 3).paragraphs:
                    for run in par.runs:
                        run.font.hidden = True
    Doc.save(File.parent.as_posix() + '/Prep_' + File.name)
